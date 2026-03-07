from docx import Document


def export_answers(questions, answers, citations):

    doc = Document()

    doc.add_heading("Completed Questionnaire", level=1)

    for i in range(len(questions)):

        doc.add_paragraph(f"Question: {questions[i]}")
        doc.add_paragraph(f"Answer: {answers[i]}")
        doc.add_paragraph(f"Citation: {citations[i]}")
        doc.add_paragraph("")

    file_path = "completed_questionnaire.docx"
    doc.save(file_path)

    return file_path
